# Import python packages
# from snowflake.snowpark import Session
import json
from dataclasses import dataclass
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import os
import streamlit as st

import tempfile
import os
import subprocess
from datetime import datetime
import io
from io import BytesIO
import zipfile
import os
import sys
import shutil
from snowflake.snowpark import Session

sys.path.append(os.path.abspath("./../../"))


def get_procedure_text(procedure_name,sharepoint_path, logger=print):
    """
    Retrieves the text of the specified stored procedure and removes snowconvert comments.
    
    Args:
        procedure_name (str): Name of the stored procedure (case-insensitive).
    
    Returns:
        str: The procedure definition text with comments removed, or '' if not found.
    """
    try:

            file_path = os.path.join(sharepoint_path, f'{procedure_name}.sql')
            
            # Check if file exists
            if not os.path.exists(file_path):
                logger(f"File not found: {file_path}")
                return ''
            
            # Read the file
            with open(file_path, 'r', encoding='utf-8') as f:
                procedure_text = f.read()
            
            logger(f"Successfully loaded: {procedure_name}.sql")
            logger(f"Generation in progress...")
        
            # More comprehensive pattern to match the entire COMMENT block
            # Matches COMMENT='{ any JSON content with "sf_sc" origin }'
            comment_pattern = r"COMMENT='\s*\{[^}]*\\\"origin\\\":\s*\\\"sf_sc\\\"[^}]*\}\s*'"
            cleaned_text = re.sub(comment_pattern, '', procedure_text, flags=re.DOTALL)
            
            # Alternative pattern if the above doesn't work - matches the exact structure
            if 'COMMENT=' in cleaned_text and 'sf_sc' in cleaned_text:
                # Fallback pattern for exact match
                exact_pattern = r"COMMENT='\s*\{\s*\\\"origin\\\":\s*\\\"sf_sc\\\",\s*\\\"name\\\":\s*\\\"snowconvert\\\",\s*\\\"version\\\":\s*\{[^}]*\},\s*\\\"attributes\\\":\s*\{[^}]*\}\s*\}\s*'"
                cleaned_text = re.sub(exact_pattern, '', cleaned_text, flags=re.DOTALL)
            
            # Clean up any extra whitespace and newlines
            cleaned_text = re.sub(r'\n\s*\n+', '\n', cleaned_text)  # Multiple blank lines to single
            cleaned_text = re.sub(r'^\s+', '', cleaned_text, flags=re.MULTILINE)  # Leading spaces
            cleaned_text = cleaned_text.strip()
            
            return cleaned_text

    except Exception as e:
        logger(f"Error retrieving text for '{procedure_name}': {e}")
        return ''

def generate_test_cases_prompt(stored_procedure_text,session, output_file_path=None,logger=print):
    """
    Generate system prompt for AI-based test case generation with stored procedure as input
    
    Args:
        stored_procedure_text (str): The complete Snowflake stored procedure code to analyze
        output_file_path (str, optional): Path to save the test cases output. If None, auto-generates filename.
    
    Returns:
        str: Generated test cases content
    """
    system_prompt = f"""

    You are an expert test engineer specializing in Snowflake stored procedure testing and validation. Your task is to analyze Snowflake SQL stored procedures and generate comprehensive test cases that achieve maximum coverage of all technical functionalities, embedded business rules, and Snowflake-specific features.
    ## Stored Procedure to Analyze:
       ```sql
        {stored_procedure_text}
        ```

    ## Your Objectives:
    1. **Identify All Technical Functionalities**: Parse the Snowflake stored procedure code to identify every distinct technical operation including Snowflake-specific functions, data types, and SQL constructs
    2. **Extract Business Rules**: Identify all conditional logic, business constraints, validation rules, data quality checks, and decision points embedded in the code
    3. **Generate Unique Test Cases (UTC)**: Create test cases that specifically target each identified functionality and business rule with Snowflake-compatible validation methods
    4. **Ensure Complete Coverage**: Iterate through each functionality to generate multiple test scenarios including positive, negative, edge cases, and Snowflake-specific scenarios

    ## Analysis Framework:
    When analyzing a Snowflake stored procedure, systematically identify EVERY:

    ### Data Operations:
    - **DML Operations**: SELECT, INSERT, UPDATE, DELETE, MERGE, COPY INTO operations
    - **JOIN Operations**: INNER JOIN, LEFT JOIN, RIGHT JOIN, FULL OUTER JOIN, CROSS JOIN- analyze each join condition and purpose
    - **Snowflake Functions**: Built-in functions (DATE_TRUNC, PARSE_JSON, FLATTEN, etc.)
    - **WHERE Clauses**: All filtering conditions, comparison operators, logical operators (AND, OR, NOT), IN/NOT IN, EXISTS/NOT EXISTS, LIKE patterns, NULL checks
    - **Window Functions**: ROW_NUMBER(), RANK(), LAG/LEAD operations
    - **Aggregate Functions**: SUM, COUNT, MIN, MAX with OVER clauses
    - **Time Travel Queries**: AT/BEFORE timestamp operations
    - **Subqueries and CTEs**: All nested queries and their conditions
    - **Set Operations**: UNION, INTERSECT, EXCEPT clauses

    ### Control Flow & Logic:
    - **Conditional Statements**: IF-ELSE, CASE WHEN constructs
    - **CASE Statement Analysis**: Identify every CASE WHEN construct and generate tests for:
      - Each WHEN condition (true and false scenarios)
      - ELSE branch execution
      - NULL handling in conditions
      - String case sensitivity in WHEN clauses
      - Numeric boundary conditions
      - Subquery results in CASE conditions
    - **Loops**: FOR loops, WHILE loops with cursors
    - **Exception Handling**: TRY-CATCH blocks, SQLSTATE handling
    - **Dynamic SQL**: EXECUTE IMMEDIATE statements
    - **Transactions**: COMMIT, ROLLBACK, transaction isolation
    - **Variable Operations**: All variable declarations, assignments, and usage
    - **Procedure Calls**: All CALL statements with parameter passing

    ### Snowflake Features:
    - **Data Types**: VARIANT, ARRAY, OBJECT handling
    - **Semi-structured Data**: JSON parsing, XML operations
    - **Clustering**: Cluster key operations and optimization
    - **Security**: Role-based access, row-level security
    - **Performance**: Query optimization, warehouse scaling

    ### Business Logic Patterns:
    - **Data Validation**: NULL checks, data type validations, range checks
    - **Data Transformations**: ETL logic, data cleansing operations
    - **Business Calculations**: Financial calculations, aggregations, derived metrics
    - **Data Quality Rules**: Duplicate detection, referential integrity
    - **Audit Trails**: Change tracking, logging mechanisms

    ## Test Case Generation Requirements:
    For EVERY identified functionality, generate ALL test scenarios:

    ### 1. Positive Test Cases:
    - Valid input parameters with expected data types
    - Normal business scenarios with realistic datasets
    - Successful execution paths through all branches
    - Valid data combinations and relationships

    ### 2. Negative Test Cases - **GENERATE EXTENSIVELY AS NEEDED FOR ROBUST TESTING**:
    - Invalid parameters (wrong data types, out of range values)
    - Missing required parameters or NULL values
    - Constraint violations and referential integrity failures
    - No-match joins: JOIN scenarios with no matching records, orphaned data
    - WHERE conditions that evaluate to false, invalid filter criteria
    - **Missing dimension table records causing failed joins**
    - **Inactive or deleted records that don't meet filter criteria (ACTIVE_IND = 'N', DELETED_IND = 'Y')**
    - **Empty source tables with zero records to process**
    - **Mismatched foreign key relationships breaking join chains**
    - **Records already up-to-date (no change scenarios)**
    - Permission errors and access control violations
    - Resource limitations and timeout scenarios
    - Invalid WHERE clause conditions and malformed queries
    
    ### 3. Boundary Test Cases:
    - Edge values for numeric ranges and date boundaries
    - Empty datasets and single-row datasets
    - Maximum/minimum values for Snowflake data types
    - Large dataset scenarios that test performance limits

    ### 4. WHERE Clause Test Cases:
   - All comparison operations (=, <>, <, >, <=, >=) with various data types
   - Logical operations (AND, OR, NOT) and operator precedence
   - Pattern matching (LIKE, REGEXP) with wildcards and special characters
   - Set operations (IN/NOT IN, EXISTS/NOT EXISTS) with various scenarios
   - NULL handling (IS NULL, IS NOT NULL) and three-valued logic
   - Range conditions (BETWEEN) with different data types

   ### 5. STRING OPERATIONS TEST CASES - **CRITICAL FOR SQL PROCEDURES**:
    - **Case Sensitivity Testing**: Test all string comparisons with various case combinations
      - Uppercase, lowercase, mixed case variations (e.g., 'CHILD', 'child', 'Child', 'ChIlD')
      - Test UPPER(), LOWER(), RTRIM(), LTRIM() function behavior with edge cases
    - **Special Characters in Strings**: 
      - Leading/trailing whitespace, tabs, newlines (' CHILD', 'CHILD ', 'CHILD\t')
      - Special characters (quotes, apostrophes, unicode)
      - Empty strings vs NULL vs whitespace-only strings
    - **String Function Edge Cases**:
      - RTRIM/LTRIM with different whitespace characters
      - String concatenation with NULL values
      - Pattern matching edge cases (LIKE operator variations)
    - **WHERE Clause String Comparisons**:
      - Test every string equality/inequality with case variations
      - Test string functions in WHERE clauses with edge inputs
      - Test UPPER(RTRIM()) combinations with various inputs

    ### 6. JOIN Test Cases:
   - INNER JOIN: Matching records, non-matching records, duplicate keys
   - LEFT JOIN: Missing right-side data, NULL propagation, multiple matches
   - RIGHT JOIN: Missing left-side data, asymmetric relationships
   - FULL OUTER JOIN: All combination scenarios, NULL handling from both sides
   - Multiple JOINs: Complex join chains and their interactions
   - Performance: Large table joins, optimization validation

   ### 7. CASE STATEMENT TEST CASES - **CRITICAL FOR CONDITIONAL LOGIC**:
    - **All CASE Branches**: Test every WHEN condition and ELSE branch
      - True conditions that trigger WHEN clauses
      - False conditions that fall through to ELSE
      - NULL handling in CASE conditions
      - Edge values that test boundary conditions in WHEN clauses
    - **Nested CASE Statements**: Test complex nested conditional logic
    - **CASE in SELECT vs UPDATE**: Test CASE statements in different contexts
    - **String Comparisons in CASE**: 
      - Case sensitivity variations in WHEN conditions (test 'CLOSED', 'closed', 'Closed')
      - UPPER/RTRIM combinations with different inputs
      - Special characters and whitespace in string conditions
    - **Numeric Comparisons in CASE**:
      - Boundary values (=, <, >, <=, >=)
      - Negative numbers, zero, positive numbers
      - NULL numeric values in conditions
    - **Subquery-based CASE Conditions**: 
      - Test when subqueries return expected values
      - Test when subqueries return NULL or no results
      - Test when subqueries return multiple results

    ### 8. Snowflake-Specific Test Cases:
    - Time zone handling and timestamp precision
    - Semi-structured data edge cases (malformed JSON, nested objects)
    - Large dataset performance testing
    - Warehouse scaling scenarios
    - Role and permission boundary testing

    ### 9. Business Rule Validation:
    - Test cases that verify specific business logic and calculations
    - Data integrity and consistency checks
    - Cross-table validation and dependency testing

    ## Test Case Prioritization:
    - High Priority: Critical business paths, essential data integrity rules, security validations, primary workflow functionality
    - Medium Priority: Secondary business rules, performance optimizations, edge case handling
    - Low Priority: Cosmetic validations, non-critical error scenarios, optional features

    ## Test Case Design Principles:
    - Independence: Each test case should be self-contained and not depend on other test execution
    - Repeatability: Test cases should produce consistent results across multiple executions
    - Minimal Setup: Use minimal, focused test data that clearly demonstrates the test scenario
    - Clear Validation: Each test should have specific, measurable success/failure criteria

    ## Output Format for Each Test Case:
    Test Case ID: TC_[FUNCTIONALITY][SCENARIO][NUMBER] 
    Functionality: [Specific technical functionality being tested] 
    Business Rule: [Associated business rule or condition] 
    Test Category: [Positive/Negative/Boundary/Integration/Performance] 
    Test Scenario: [Detailed description of what is being tested] 
    Preconditions: [Required setup, data state, or permissions] 
    Test Data Setup: [CREATE TABLE statements, INSERT statements for test data] 
    Input Parameters: [Stored procedure parameters with values] 
    Execution Method: [CALL procedure_name(parameters) or test framework] 
    Expected Result: [Anticipated outcome, return values, data state changes] 
    Validation Queries: [SELECT statements to verify results OR error validation for negative cases] 
    Cleanup: [Statements to clean up test data] 
    Priority: [High/Medium/Low based on business impact]

    ## Validation Methods:
    - **Result Set Validation**: Compare actual vs expected query results
    - **Data State Verification**: Check table contents before/after execution
    - **Return Code Validation**: Verify stored procedure return values
    - **Exception Testing**: Validate error messages and SQLSTATE codes
    - **Performance Validation**: Execution time and resource usage checks
    - **Audit Log Verification**: Check system logs and audit trails

    ## Coverage Requirements - MANDATORY:
    Ensure test cases cover:
    - Statement Coverage: Every SQL statement executed at least once
    - Branch Coverage: Every conditional path tested (IF-ELSE, CASE branches)
    - WHERE Clause Coverage: Every filtering condition tested with multiple scenarios
    - JOIN Coverage: Every JOIN operation tested with multiple data scenarios including NULL handling
    - Function Coverage: Every built-in and user-defined function tested
    - Business Rule Coverage: Every business logic condition validated
    - Data Coverage: Different data scenarios and edge cases
    - Error Coverage: All exception handling paths tested with proper negative scenarios
    - Integration Coverage: End-to-end business process validation
    - String Operations Coverage: All string comparisons tested with case sensitivity and special character variations
    - CASE Statement Coverage: Every WHEN and ELSE branch tested with comprehensive scenarios
    - Negative Scenario Coverage: Generate comprehensive negative test cases to cover all possible failure paths and edge conditions

    Generate executable, specific test cases with clear setup instructions, validation criteria, and cleanup procedures for comprehensive Snowflake stored procedure testing  based on the provided stored procedure code above.
    ENSURE COMPLETE COVERAGE of every functionality identified
    Follow the test case generation guidelines to create independent, repeatable test cases with minimal data setup and clear prioritization.

    **CRITICAL FORMATTING REQUIREMENTS FOR SQL CODE IDENTIFICATION**:

    1. **SQL Code Block Markers**: 
    - Use [SQL CODE BLOCK START] and [SQL CODE BLOCK END] to wrap ALL SQL statements
    - This includes: CREATE, INSERT, UPDATE, DELETE, SELECT, CALL, DROP, TRUNCATE, MERGE, etc.

    2. **What to Mark as SQL Code**:
    - Table creation statements: CREATE TABLE, CREATE TEMP TABLE
    - Data manipulation: INSERT INTO, UPDATE SET, DELETE FROM, MERGE
    - Query statements: SELECT FROM, WITH clauses, CTEs
    - Procedure calls: CALL procedure_name(parameters)
    - Cleanup statements: DROP TABLE, TRUNCATE TABLE
    - Transaction control: BEGIN, COMMIT, ROLLBACK
    - Variable declarations: DECLARE, SET
    - Control flow in SQL: IF-ELSE blocks, WHILE loops, FOR loops
    - Exception handling: TRY-CATCH blocks
    - Any multi-line SQL that spans multiple lines

    3. **What to Leave as Regular Text**:
    - Test case descriptions and explanations
    - Field labels (Functionality:, Test Scenario:, Expected Result:)
    - Plain text explanations of what the test does
    - Priority levels (High/Medium/Low)
    - Test categories (Positive/Negative/Boundary)
    - Expected outcomes described in words
    - Error messages or SQLSTATE codes when written as text descriptions

    4. **Examples of Proper Marking**:

    Test Data Setup: 
    [SQL CODE BLOCK START]
    CREATE TEMP TABLE test_customers (
        customer_id INT,
        customer_name VARCHAR(100),
        status VARCHAR(20)
    );
    INSERT INTO test_customers VALUES 
    (1, 'John Doe', 'Active'),
    (2, 'Jane Smith', 'Inactive');
    [SQL CODE BLOCK END]

    Execution Method:
    [SQL CODE BLOCK START]
    CALL DIM_CRM_BUS_UNIT_BUILD();
    [SQL CODE BLOCK END]

    Expected Result: The procedure should complete successfully and return 'SUCCESS' message.

    Validation Queries:
    [SQL CODE BLOCK START]
    SELECT COUNT(*) FROM target_table WHERE status = 'Processed';
    SELECT * FROM audit_log WHERE procedure_name = 'DIM_CRM_BUS_UNIT_BUILD';
    [SQL CODE BLOCK END]

    5. **Single Line SQL**: Even single SQL statements should be marked:
    [SQL CODE BLOCK START]
    SELECT COUNT(*) FROM test_table;
    [SQL CODE BLOCK END]

    6. **Multiple SQL Statements**: Group related SQL statements in one block when they work together:
    [SQL CODE BLOCK START]
    BEGIN TRANSACTION;
    INSERT INTO staging_table SELECT * FROM source_table;
    UPDATE staging_table SET processed_date = CURRENT_TIMESTAMP();
    COMMIT;
    [SQL CODE BLOCK END]

    **IMPORTANT**: Be consistent - if it's executable SQL code, wrap it in markers. If it's descriptive text about the test or expected behavior, leave it as regular text.

    """

    prompt = system_prompt.replace("'", "''")
 
    # Call Cortex COMPLETE via SQL
    cortex_query = f"""
    SELECT SNOWFLAKE.CORTEX.COMPLETE(
        'claude-4-sonnet',
        [
            {{
                'role': 'user',
                'content': '{prompt}'
            }}
        ],
        {{
            'temperature': 0.1,
            'max_tokens': 32000
        }}
    )
    """
    
    try:
        cortex_df = session.sql(cortex_query)
        response_json = cortex_df.collect()[0][0]
        
        # Parse the JSON response correctly
        response_data = json.loads(response_json)
        test_cases_content = response_data['choices'][0]['messages'].strip()
        
        logger("Test cases generated successfully!")
        return test_cases_content
        
    except Exception as e:
        error_msg = f"Error generating test cases: {str(e)}"
        logger(error_msg)
        return error_msg

def create_professional_docx(stored_procedure_text, test_cases_content, output_folder, procedure_name=None, output_file_path=None, logger=print):
    """
    Create a professional Word document with formatting, colors, and styles
    """
    try:

        # Generate output file path if not provided
        if output_file_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Create testcases folder in current working directory
            #output_folder = "testcases/SP/Done"
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
                print(f"Created folder: {output_folder}")
            
            if procedure_name:
                filename = f"{procedure_name}_test_cases_{timestamp}.docx"
            else:
                filename = f"test_cases_output_{timestamp}.docx"
            
            output_file_path = os.path.join(output_folder, filename)
        else:
            # If custom path provided, ensure the directory exists
            output_dir = os.path.dirname(output_file_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                logger(f"Created folder: {output_dir}")
        
        doc = Document()
        
        # Define custom styles (keeping existing styles...)
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 70, 140)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(20)
        
        h1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Calibri'
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(31, 78, 120)
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
        
        h2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Calibri'
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(68, 114, 196)
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(4)
        
        tc_style = doc.styles.add_style('TestCase', WD_STYLE_TYPE.PARAGRAPH)
        tc_font = tc_style.font
        tc_font.name = 'Calibri'
        tc_font.size = Pt(12)
        tc_font.bold = True
        tc_font.color.rgb = RGBColor(112, 48, 160)
        tc_style.paragraph_format.space_before = Pt(8)
        tc_style.paragraph_format.space_after = Pt(4)
        
        field_style = doc.styles.add_style('FieldLabel', WD_STYLE_TYPE.CHARACTER)
        field_font = field_style.font
        field_font.name = 'Calibri'
        field_font.size = Pt(11)
        field_font.bold = True
        field_font.color.rgb = RGBColor(89, 89, 89)
        
        normal_style = doc.styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = 'Calibri'
        normal_font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(3)
        
        # Helper function to add code paragraph with background color
        def add_code_with_background(text_content):
            from docx.oxml.shared import OxmlElement, qn
            
            para = doc.add_paragraph()
            run = para.add_run(text_content)
            
            # Set font properties
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(25, 25, 112)  # Dark blue text
            
            # Add light gray background
            try:
                shading_elm = OxmlElement(qn('w:shd'))
                shading_elm.set(qn('w:fill'), 'F0F0F0')  # Light gray background
                run._element.get_or_add_rPr().append(shading_elm)
            except:
                # If background fails, just use the text formatting
                pass
            
            # Set paragraph formatting
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.space_before = Pt(2)
            
            return para
        
        # AGGRESSIVE MARKDOWN CLEANUP FUNCTION
        def clean_markdown_completely(content):
            """Remove ALL markdown formatting aggressively"""
            
            # Remove markdown headers (##, ###, ####, etc.)
            content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
            content = re.sub(r'^#{1,6}([^#\n]+)#{1,6}$', r'\1', content, flags=re.MULTILINE)
            
            # Remove bold/italic markers
            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)  # **bold**
            #content = re.sub(r'\*([^*]+)\*', r'\1', content)      # *italic*
            content = re.sub(r'__([^_]+)__', r'\1', content)      # __bold__
            #content = re.sub(r'_([^_]+)_', r'\1', content)        # _italic_
            
            # Remove code backticks but preserve SQL markers
            content = re.sub(r'`([^`]+)`', r'\1', content)        # `inline code`
            content = re.sub(r'```[^\n]*\n(.*?)\n```', r'\1', content, flags=re.DOTALL)  # ```code blocks```
            
            # Remove list markers
            content = re.sub(r'^\s*[-*+]\s+', '', content, flags=re.MULTILINE)  # - * + lists
            content = re.sub(r'^\s*\d+\.\s+', '', content, flags=re.MULTILINE)  # 1. 2. numbered lists
            
            # Remove horizontal rules
            content = re.sub(r'^-{3,}$', '', content, flags=re.MULTILINE)
            content = re.sub(r'^\*{3,}$', '', content, flags=re.MULTILINE)
            
            # Remove extra whitespace
            content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)  # Multiple blank lines
            content = re.sub(r'^\s+', '', content, flags=re.MULTILINE)  # Leading spaces
            content = re.sub(r'\s+$', '', content, flags=re.MULTILINE)  # Trailing spaces
            
            return content.strip()
        
        # Add document title
        doc.add_paragraph('AI-Generated Test Cases for Snowflake Stored Procedure', style='CustomTitle')
        
        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run('Generated on: ').font.bold = True
        meta_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Add original stored procedure section
        doc.add_paragraph('Original Stored Procedure', style='CustomHeading1')
        add_code_with_background(stored_procedure_text)
        
        # Add generated test cases section
        doc.add_paragraph('Generated Test Cases', style='CustomHeading1')
        
        # CLEAN THE CONTENT FIRST - REMOVE ALL MARKDOWN
        cleaned_content = clean_markdown_completely(test_cases_content)
        
        logger("Markdown cleanup completed. Sample of cleaned content:")
        logger(cleaned_content[:500] + "..." if len(cleaned_content) > 500 else cleaned_content)
        
        # Process test cases content with SQL markers
        def process_content_with_markers(content):
            lines = content.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i].strip()
                
                # Skip empty lines
                if not line:
                    i += 1
                    continue
                
                # Check for SQL code block START marker
                if '[SQL CODE BLOCK START]' in line:
                    # Skip the START marker line itself
                    i += 1
                    sql_block = []
                    
                    # Collect SQL lines until END marker
                    while i < len(lines):
                        current_line = lines[i]
                        if '[SQL CODE BLOCK END]' in current_line:
                            # Found END marker, stop collecting and skip this line
                            i += 1
                            break
                        else:
                            # Add this line to SQL block
                            sql_block.append(current_line)
                            i += 1
                    
                    # Add SQL block to document with background (without the markers)
                    if sql_block:
                        # Clean up the SQL block
                        while sql_block and not sql_block[0].strip():
                            sql_block.pop(0)
                        while sql_block and not sql_block[-1].strip():
                            sql_block.pop()
                        
                        if sql_block:
                            add_code_with_background('\n'.join(sql_block))
                
                # Check for main headings (now without ## markers)
                elif any(keyword in line.lower() for keyword in [
                    'test plan', 'executive summary', 'technical functionality', 
                    'test execution', 'coverage analysis', 'comprehensive test plan',
                    'analysis framework', 'technical functionalities'
                ]):
                    doc.add_paragraph(line, style='CustomHeading1')
                    i += 1
                
                # Check for sub headings (now without ## markers)
                elif any(keyword in line.lower() for keyword in [
                    'core operations', 'business rules', 'functionality analysis',
                    'phase 1', 'phase 2', 'phase 3', 'test cases for',
                    'data operations', 'control flow', 'snowflake features',
                    'business logic patterns', 'test case generation'
                ]):
                    doc.add_paragraph(line, style='CustomHeading2')
                    i += 1
                
                # Check for test case IDs
                elif (line.startswith('Test Case ID:') or line.startswith('TC_') or 
                      'Test Case ID:' in line):
                    doc.add_paragraph(line, style='TestCase')
                    i += 1
                
                # Check for field labels
                elif ':' in line and any(keyword in line for keyword in [
                    'Functionality', 'Business Rule', 'Test Category', 'Test Scenario',
                    'Preconditions', 'Expected Result', 'Priority', 'Input Parameters',
                    'Execution Method', 'Validation Queries', 'Cleanup', 'Test Data Setup'
                ]):
                    para = doc.add_paragraph()
                    parts = line.split(':', 1)
                    if len(parts) == 2:
                        label_run = para.add_run(parts[0] + ': ')
                        label_run.font.name = 'Calibri'
                        label_run.font.size = Pt(11)
                        label_run.font.bold = True
                        label_run.font.color.rgb = RGBColor(89, 89, 89)
                        
                        content_run = para.add_run(parts[1].strip())
                        content_run.font.name = 'Calibri'
                        content_run.font.size = Pt(11)
                    else:
                        para.add_run(line)
                    i += 1
                
                # Regular paragraphs (skip any remaining markers)
                elif not ('[SQL CODE BLOCK' in line):
                    para = doc.add_paragraph(line, style='NormalText')
                    i += 1
                else:
                    # Skip any remaining marker lines
                    i += 1
        
        # Process the cleaned content
        process_content_with_markers(cleaned_content)
        
        # Save document
        doc.save(output_file_path)
        
        logger(f"Professional test cases document successfully saved to: {output_file_path}")
        logger("Document features:")
        logger("- ALL markdown formatting completely removed")
        logger("- SQL code blocks with dark blue text and light gray background")
        logger("- SQL markers properly removed")
        logger("- Professional styling and layout")
        
        return output_file_path
        
    except Exception as e:
        error_msg = f"Error creating DOCX document: {str(e)}"
        logger(error_msg)
        return None
def process_all_sql_files(sharepoint_path,output_folder,session, logger=print):
    """
    Simple loop to process all SQL files in the SharePoint folder
    """
    import os
    
    # Get all SQL files
    files = os.listdir(sharepoint_path)
    sql_files = [f.replace('.sql', '') for f in files if f.endswith('.sql')]
    
    logger(f"Processing {len(sql_files)} SQL files...")
    
    # Simple loop
    for procedure_name in sql_files:
        logger(f"Processing: {procedure_name}")
        
        # Get procedure text
        stored_proc = get_procedure_text(procedure_name,sharepoint_path)
        
        # Generate test cases
        test_cases_content = generate_test_cases_prompt(stored_proc,session)
        
        # Create DOCX
        create_professional_docx(stored_proc, test_cases_content, output_folder, procedure_name)
        
        logger(f"Completed: {procedure_name}")


# process_all_sql_files(sharepoint_path,output_folder)

def test_cases(sf_session):
    
    # st.markdown("<h1>🚀 SP Test Cases</h1>", unsafe_allow_html=True)
    st.markdown("<h2>🗒️ SP Test Cases Generation</h2>", unsafe_allow_html=True)
    # st.markdown("---")
    st.markdown(
        """Generating Testcases for the snowflake Stored Procedures."""
    )
    session=sf_session


    if "sp_logs" not in st.session_state:
        st.session_state.sp_logs = []
    if "sp_done" not in st.session_state:
        st.session_state.sp_done = False
    if "sp_zip_output" not in st.session_state:
        st.session_state.sp_zip_output = None


    with st.expander("📦 Upload ZIP File", expanded=True):
        input_folder = st.file_uploader("Upload a ZIP folder containing SQLs", type="zip")
        # output_folder = st.text_input("🔹 Enter Output Folder Path", placeholder="e.g., testcases/SP/Deployed")


    def streamlit_logger(msg: str):
        st.session_state.sp_logs.append(msg)
        if not st.session_state.sp_done:
            log_placeholder.text_area("📜 Conversion Logs", "\n".join(st.session_state.sp_logs), height=300, )


    def zip_folder(folder_path):
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, _, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, folder_path)
                    zipf.write(file_path, arcname)
        zip_buffer.seek(0)
        return zip_buffer

    if not input_folder:
        st.session_state.sp_logs = []

    if input_folder:
        if st.button("🚀 Run Test Case", use_container_width=True):
            st.session_state.sp_logs = []
            st.session_state.sp_done = False
            st.session_state.sp_zip_output = None
            log_placeholder = st.empty()

            with tempfile.TemporaryDirectory() as temp_dir:
                zip_path = os.path.join(temp_dir, "uploaded.zip")
                with open(zip_path, "wb") as f:
                    f.write(input_folder.read())

                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)

                source_sql_dir = os.path.join(temp_dir, "source_sqls")
                output_sql_dir = os.path.join(temp_dir, "transformed_sqls")
                os.makedirs(source_sql_dir, exist_ok=True)
                os.makedirs(output_sql_dir, exist_ok=True)

                # Move all .sql files to source_sql_dir
                for root, _, files in os.walk(temp_dir):
                    for file in files:
                        if file.endswith(".sql"):
                            src_path = os.path.join(root, file)
                            dst_path = os.path.join(source_sql_dir, file)
                            if os.path.abspath(src_path) != os.path.abspath(dst_path):
                                shutil.copy(src_path, dst_path)

                # Run transformation
                
                streamlit_logger("🔄 Starting SP Test Case Generation...")
                with st.spinner('Generating Test Cases... please wait ⏳'):
                    process_all_sql_files(source_sql_dir, output_sql_dir, session,logger=streamlit_logger)
                streamlit_logger("✅ Test Case Generation Completed.")

                # Zip output
                zipped_output = zip_folder(output_sql_dir)
                st.session_state.sp_zip_output = zipped_output
                st.session_state.sp_done = True

                log_placeholder.empty()

    # ----------------------------
    # 📜 Show Logs Before Download
    # ----------------------------
    if st.session_state.sp_done and st.session_state.sp_logs:
        st.text_area("📜 Conversion Logs", "\n".join(st.session_state.sp_logs), height=300, key="sp_static_logs")

    # ----------------------------
    # 📥 Download Button
    # ----------------------------
    if st.session_state.sp_done and st.session_state.sp_logs and st.session_state.sp_zip_output:
        st.download_button(
            label="📄 Download Testcases",
            data=st.session_state.sp_zip_output,
            file_name="Test_cases_generated.zip",
            mime="application/zip",
            key="download_test_cases"
        )

